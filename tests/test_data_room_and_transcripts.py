"""Wiring tests for C2 (data room) + C3 (transcripts).

Covers:
  - parsers (txt/csv directly; pdf/xlsx/docx via real-shaped fixtures)
  - store roundtrip for both
  - agent tools query_data_room / query_transcripts
  - API endpoints (upload, list, delete)
"""
from __future__ import annotations

import io

import pytest
from fastapi.testclient import TestClient


# ---------------------------------------------------------------------------
# Parsers
# ---------------------------------------------------------------------------


def test_data_room_parse_txt(tmp_path):
    from marvin.ingestion.data_room import parse_file
    p = tmp_path / "notes.txt"
    p.write_text("Hello\nWorld\nNet revenue 123\n")
    out = parse_file(p)
    assert "Net revenue 123" in out


def test_data_room_parse_csv(tmp_path):
    from marvin.ingestion.data_room import parse_file
    p = tmp_path / "data.csv"
    p.write_text("year,revenue\n2024,100\n2025,150\n")
    out = parse_file(p)
    assert "2024" in out and "150" in out


def test_data_room_parse_xlsx_roundtrip(tmp_path):
    from openpyxl import Workbook
    from marvin.ingestion.data_room import parse_file

    wb = Workbook()
    ws = wb.active
    ws.title = "P&L"
    ws.append(["year", "revenue", "ebitda"])
    ws.append([2024, 100.5, 22.0])
    p = tmp_path / "pl.xlsx"
    wb.save(str(p))
    out = parse_file(p)
    assert "[sheet: P&L]" in out
    assert "ebitda" in out
    assert "100.5" in out


def test_data_room_parse_docx_roundtrip(tmp_path):
    import docx
    from marvin.ingestion.data_room import parse_file

    d = docx.Document()
    d.add_paragraph("Mgmt thesis")
    d.add_paragraph("Net retention exceeds 120%")
    p = tmp_path / "memo.docx"
    d.save(str(p))
    out = parse_file(p)
    assert "Mgmt thesis" in out
    assert "Net retention exceeds 120%" in out


def test_data_room_parse_unsupported(tmp_path):
    from marvin.ingestion.data_room import parse_file
    p = tmp_path / "foo.zip"
    p.write_bytes(b"junk")
    with pytest.raises(ValueError, match="unsupported"):
        parse_file(p)


def test_search_text_returns_snippets():
    from marvin.ingestion.data_room import search_text
    text = "alpha\nbeta gamma DELTA epsilon\nzeta delta\n"
    hits = search_text(text, "delta")
    assert len(hits) == 2
    assert hits[0]["line"] == 2
    assert hits[1]["line"] == 3


# ---------------------------------------------------------------------------
# Transcript parser
# ---------------------------------------------------------------------------


def test_parse_transcript_q_a_format():
    from marvin.ingestion.transcripts import parse_transcript
    raw = (
        "Q: What is your net retention?\n"
        "A: Last quarter we ran at 125 percent, down from 130 a year ago.\n"
        "Q: And gross margin?\n"
        "A: 75 percent on a non-GAAP basis.\n"
    )
    segs = parse_transcript(raw)
    speakers = [s.speaker for s in segs]
    assert speakers == ["Q", "A", "Q", "A"]
    assert "125 percent" in segs[1].text


def test_parse_transcript_named_speakers_and_timestamps():
    from marvin.ingestion.transcripts import parse_transcript
    raw = (
        "[00:00:10] John Smith: We grew at 40% last year.\n"
        "[00:00:25] Jane Doe: Driven mostly by enterprise expansion.\n"
    )
    segs = parse_transcript(raw)
    assert segs[0].speaker == "John Smith"
    assert segs[1].speaker == "Jane Doe"
    assert "40%" in segs[0].text


def test_search_transcript_returns_segments():
    from marvin.ingestion.transcripts import parse_transcript, search_transcript
    raw = "Q: What about churn?\nA: It's around 8% gross.\n"
    segs = parse_transcript(raw)
    hits = search_transcript(segs, "churn")
    assert len(hits) == 1
    assert hits[0]["speaker"] == "Q"


# ---------------------------------------------------------------------------
# API endpoints (TestClient)
# ---------------------------------------------------------------------------


def _setup(tmp_path, monkeypatch):
    monkeypatch.setenv("MARVIN_DB_PATH", str(tmp_path / "marvin.db"))
    import importlib
    import marvin.mission.store as store_mod
    importlib.reload(store_mod)
    import marvin_ui.server as srv
    importlib.reload(srv)
    from marvin.mission.schema import Mission as MissionModel
    store = store_mod.MissionStore(str(tmp_path / "marvin.db"))
    store.save_mission(MissionModel(id="m-x", client="c", target="t", mission_type="cdd"))
    return TestClient(srv.app), store


def test_data_room_upload_list_delete(tmp_path, monkeypatch):
    client, store = _setup(tmp_path, monkeypatch)
    files = {"file": ("notes.txt", io.BytesIO(b"net retention 125 percent\n"), "text/plain")}
    r = client.post("/api/v1/missions/m-x/data-room/upload", files=files)
    assert r.status_code == 200, r.text
    fid = r.json()["id"]
    assert r.json()["parsed_text"].startswith("net retention")

    r = client.get("/api/v1/missions/m-x/data-room")
    assert len(r.json()["files"]) == 1
    assert r.json()["files"][0]["id"] == fid

    r = client.delete(f"/api/v1/missions/m-x/data-room/{fid}")
    assert r.status_code == 200
    assert r.json()["deleted"] == fid

    r = client.get("/api/v1/missions/m-x/data-room")
    assert r.json()["files"] == []


def test_transcript_upload_list_delete(tmp_path, monkeypatch):
    client, store = _setup(tmp_path, monkeypatch)
    body = {
        "text": "Q: Churn?\nA: 8% gross.\n",
        "title": "Tegus call 2025",
        "expert_name": "John Smith",
    }
    r = client.post("/api/v1/missions/m-x/transcripts", data=body)
    assert r.status_code == 200, r.text
    tid = r.json()["transcript"]["id"]
    assert r.json()["segment_count"] == 2

    r = client.get("/api/v1/missions/m-x/transcripts")
    assert len(r.json()["transcripts"]) == 1

    r = client.delete(f"/api/v1/missions/m-x/transcripts/{tid}")
    assert r.json()["deleted"] == tid


def test_query_data_room_tool_returns_hits(tmp_path, monkeypatch):
    monkeypatch.setenv("MARVIN_DB_PATH", str(tmp_path / "marvin.db"))
    import importlib
    import marvin.mission.store as store_mod
    importlib.reload(store_mod)
    import marvin.tools.data_room_tools as drt
    importlib.reload(drt)
    from marvin.mission.schema import DataRoomFile, Mission as MissionModel

    store = store_mod.MissionStore(str(tmp_path / "marvin.db"))
    store.save_mission(MissionModel(id="m-x", client="c", target="t", mission_type="cdd"))
    store.save_data_room_file(
        DataRoomFile(
            id="dr-1", mission_id="m-x", filename="notes.txt",
            file_path=str(tmp_path / "notes.txt"),
            mime_type="text/plain", size_bytes=42,
            parsed_text="line one\nnet retention 125\nline three\n",
        )
    )

    out = drt.query_data_room("net retention", state={"mission_id": "m-x"})
    assert out["files_searched"] == 1
    assert len(out["hits"]) == 1
    assert out["hits"][0]["filename"] == "notes.txt"
    assert "data_room://dr-1" in out["hits"][0]["ref"]


def test_query_transcripts_tool_returns_hits(tmp_path, monkeypatch):
    monkeypatch.setenv("MARVIN_DB_PATH", str(tmp_path / "marvin.db"))
    import importlib
    import marvin.mission.store as store_mod
    importlib.reload(store_mod)
    import marvin.tools.data_room_tools as drt
    importlib.reload(drt)
    from marvin.mission.schema import (
        Mission as MissionModel, Transcript, TranscriptSegment,
    )

    store = store_mod.MissionStore(str(tmp_path / "marvin.db"))
    store.save_mission(MissionModel(id="m-x", client="c", target="t", mission_type="cdd"))
    store.save_transcript(
        Transcript(
            id="tx-1", mission_id="m-x", title="Tegus call",
            expert_name="J. Smith", raw_text="Q: ...\nA: 125% NRR.\n",
            line_count=2,
        ),
        [
            TranscriptSegment(id="tx-1-seg-0", transcript_id="tx-1",
                              speaker="Q", text="What is NRR?", line_start=1, line_end=1),
            TranscriptSegment(id="tx-1-seg-1", transcript_id="tx-1",
                              speaker="A", text="125% NRR last quarter.", line_start=2, line_end=2),
        ],
    )

    out = drt.query_transcripts("NRR", state={"mission_id": "m-x"})
    assert out["transcripts_searched"] == 1
    assert any(h["speaker"] == "A" for h in out["hits"])
    assert any("transcript://tx-1" in h["ref"] for h in out["hits"])
