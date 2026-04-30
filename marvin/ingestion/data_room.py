"""Data-room file parsing.

MVP: PDF (pypdf), XLSX (openpyxl), DOCX (python-docx), plain text/CSV.
Returns extracted text or raises ValueError with a parse_error string.
The caller persists `parse_error` and `parsed_text=None` so partial
ingestion is visible rather than silent.
"""
from __future__ import annotations

import csv
import io
from pathlib import Path

MAX_BYTES = 25 * 1024 * 1024  # 25 MB upload limit
SUPPORTED_EXTS = {".pdf", ".xlsx", ".xls", ".docx", ".txt", ".md", ".csv"}


def parse_file(path: str | Path) -> str:
    """Parse a data-room file and return extracted text.

    Raises ValueError on unsupported extension or parse failure.
    """
    p = Path(path)
    if not p.is_file():
        raise ValueError(f"file not found: {path}")
    ext = p.suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"unsupported extension: {ext}")
    if p.stat().st_size > MAX_BYTES:
        raise ValueError(f"file exceeds {MAX_BYTES} bytes")

    if ext == ".pdf":
        return _parse_pdf(p)
    if ext in {".xlsx", ".xls"}:
        return _parse_xlsx(p)
    if ext == ".docx":
        return _parse_docx(p)
    if ext == ".csv":
        return _parse_csv(p)
    return p.read_text(encoding="utf-8", errors="replace")


def _parse_pdf(p: Path) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(str(p))
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"pdf parse failed: {exc}") from exc
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            parts.append(f"[page {i + 1}]\n{page.extract_text() or ''}")
        except Exception as exc:  # noqa: BLE001
            parts.append(f"[page {i + 1} parse error: {exc}]")
    return "\n\n".join(parts).strip()


def _parse_xlsx(p: Path) -> str:
    from openpyxl import load_workbook

    try:
        wb = load_workbook(str(p), data_only=True, read_only=True)
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"xlsx parse failed: {exc}") from exc
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"[sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(cells):
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def _parse_docx(p: Path) -> str:
    import docx

    try:
        d = docx.Document(str(p))
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"docx parse failed: {exc}") from exc
    paras = [para.text for para in d.paragraphs if para.text and para.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            if any(c.strip() for c in cells):
                paras.append("\t".join(cells))
    return "\n".join(paras).strip()


def _parse_csv(p: Path) -> str:
    text = p.read_text(encoding="utf-8", errors="replace")
    rows = list(csv.reader(io.StringIO(text)))
    return "\n".join("\t".join(r) for r in rows).strip()


def search_text(text: str, query: str, snippet_chars: int = 240, max_hits: int = 5) -> list[dict]:
    """Substring search returning [{line, snippet}, ...] for query in text.

    Case-insensitive. Each hit is a snippet centered on the match. Used by
    the agent-facing query_data_room tool — no fuzzy / semantic matching
    in the MVP.
    """
    if not query.strip():
        return []
    needle = query.lower()
    text_lower = text.lower()
    hits: list[dict] = []
    pos = 0
    while True:
        i = text_lower.find(needle, pos)
        if i < 0 or len(hits) >= max_hits:
            break
        start = max(0, i - snippet_chars // 2)
        end = min(len(text), i + snippet_chars // 2)
        snippet = text[start:end].replace("\n", " ")
        line_no = text.count("\n", 0, i) + 1
        hits.append({"line": line_no, "snippet": snippet, "char_offset": i})
        pos = i + len(needle)
    return hits
